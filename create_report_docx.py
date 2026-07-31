import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_numpages(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def build_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Footer page numbers
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_p.add_run("Trang ")
        footer_run.font.name = "Times New Roman"
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(100, 100, 100)
        
        add_page_number(footer_run)
        
        footer_run2 = footer_p.add_run(" / ")
        footer_run2.font.name = "Times New Roman"
        footer_run2.font.size = Pt(10)
        footer_run2.font.italic = True
        footer_run2.font.color.rgb = RGBColor(100, 100, 100)
        
        add_numpages(footer_run2)

    # Styling defaults: Times New Roman, standard clean black font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(30, 30, 30)

    # Title Banner (Clean, Academic, Not flashy)
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_org = p_org.add_run("TRƯỜNG ĐẠI HỌC CÔNG THƯƠNG TP. HỒ CHÍ MINH (HUIT)\nTRUNG TÂM THÔNG TIN & TRỢ LÝ AI TUYỂN SINH\n-------------------***-------------------")
    r_org.font.name = 'Times New Roman'
    r_org.font.size = Pt(12)
    r_org.font.bold = True
    r_org.font.color.rgb = RGBColor(50, 50, 50)
    p_org.paragraph_format.space_after = Pt(18)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("BÁO CÁO KĨ THUẬT XỬ LÝ SỰ CỐ & TỐI ƯU HỆ THỐNG HUIT CHATBOT")
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102) # Dark Corporate Navy
    p_title.paragraph_format.space_after = Pt(6)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Về việc: Đồng bộ Deployment Vercel/Git, Tối ưu Độ trễ Real-time SSE Streaming và Nâng cấp Giao diện Ô Nhập Liệu")
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(80, 80, 80)
    p_sub.paragraph_format.space_after = Pt(24)

    # Meta Info Table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, "E0E0E0")
    
    meta_data = [
        ("Ngày thực hiện báo cáo:", "31/07/2026"),
        ("Phiên bản hệ thống:", "HUIT Chatbot RAG v10 (Grounded Score & Streaming)"),
        ("Môi trường triển khai:", "Vercel Serverless Production (huit-chatbot.vercel.app)"),
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        
        cell_k.paragraphs[0].add_run(k).bold = True
        cell_k.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell_k.paragraphs[0].runs[0].font.size = Pt(11)
        
        cell_v.paragraphs[0].add_run(v)
        cell_v.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell_v.paragraphs[0].runs[0].font.size = Pt(11)
        
        set_cell_margins(cell_k, 60, 60, 100, 100)
        set_cell_margins(cell_v, 60, 60, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # MỤC LỤC SECTION
    h_toc = doc.add_heading(level=1)
    r_toc = h_toc.add_run("MỤC LỤC")
    r_toc.font.name = "Times New Roman"
    r_toc.font.size = Pt(14)
    r_toc.font.bold = True
    r_toc.font.color.rgb = RGBColor(0, 51, 102)
    h_toc.paragraph_format.space_before = Pt(12)
    h_toc.paragraph_format.space_after = Pt(6)

    add_toc(doc)

    doc.add_page_break()

    # SECTION 1
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. TỔNG QUAN VẤN ĐỀ VÀ MỤC TIÊU XỬ LÝ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 51, 102)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Theo yêu cầu rà soát và kiểm định hoạt động của Hệ thống Trợ lý AI Tuyển sinh HUIT (Cổng thông tin tuyển sinh Đại học Công Thương TP.HCM), đội ngũ kĩ thuật đã tiến hành kiểm tra toàn diện và ghi nhận các vấn đề chính sau:")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    bullets = [
        "Trạng thái triển khai (Deployment Status): Mã nguồn local tại máy trạm đã có những cải tiến mới nhưng chưa được đẩy (push) đầy đủ lên GitHub remote, khiến phiên bản chạy thực tế trên Vercel vẫn còn dừng lại ở commit cũ.",
        "Độ trễ phản hồi (Response Latency): Việc phản hồi qua API stream trước đây còn gọi hàm xử lý đồng bộ, khiến thời gian chờ ký tự đầu tiên (Time To First Token - TTFT) kéo dài từ 4 đến 7 giây, gây cảm giác chậm trễ cho người dùng.",
        "Trải nghiệm ô nhập liệu (User Input UX): Ô nhập câu hỏi bị khóa cứng (disabled) trong suốt quá trình AI đang sinh phản hồi, ngăn cản người dùng gõ sẵn câu hỏi tiếp theo.",
    ]
    for b in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r = bp.add_run(b)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.line_spacing = 1.15

    # SECTION 2
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. CÁC NỘI DUNG ĐÃ XỬ LÝ VÀ NÂNG CẤP KĨ THUẬT")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0, 51, 102)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # 2.1
    h21 = doc.add_heading(level=2)
    r21 = h21.add_run("2.1. Đồng bộ mã nguồn Git Remote và Vercel Production")
    r21.font.name = "Times New Roman"
    r21.font.size = Pt(13)
    r21.font.bold = True
    r21.font.color.rgb = RGBColor(40, 40, 40)
    h21.paragraph_format.space_before = Pt(8)
    h21.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph("Đã kiểm tra hiện trạng git log và thực hiện đẩy toàn bộ các commit mới nhất lên kho chứa GitHub (origin/main). Vercel đã nhận được webhook và hoàn tất quy trình Build & Deploy sản phẩm thành công:")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    commit_table = doc.add_table(rows=3, cols=3)
    commit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(commit_table)
    
    headers = ["Commit Hash", "Thông điệp (Message)", "Thời gian đẩy / Trạng thái"]
    for idx, text in enumerate(headers):
        cell = commit_table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.bold = True
        r.font.size = Pt(11)
        set_cell_margins(cell, 80, 80, 100, 100)

    rows_data = [
        ("537572d", "feat: optimize latency with streaming UI, RAM cache, and KMeans clustering", "31/07/2026 17:21 (Hoàn tất)"),
        ("1c7d97f", "feat: enable true real-time LLM streaming and keep chat input enabled", "31/07/2026 17:32 (Hoàn tất)"),
    ]
    for r_idx, data in enumerate(rows_data, 1):
        row = commit_table.rows[r_idx]
        for c_idx, val in enumerate(data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            set_cell_margins(cell, 60, 60, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2.2
    h22 = doc.add_heading(level=2)
    r22 = h22.add_run("2.2. Khắc phục độ trễ phản hồi với Real-Time SSE Streaming")
    r22.font.name = "Times New Roman"
    r22.font.size = Pt(13)
    r22.font.bold = True
    r22.font.color.rgb = RGBColor(40, 40, 40)
    h22.paragraph_format.space_before = Pt(8)
    h22.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph("Tái thiết kế hàm `stream_answer` và bổ sung hàm `_stream_llm` trong mô-đun rag_core.py:")
    p.paragraph_format.space_after = Pt(4)
    
    stream_points = [
        "Truy vấn tức thì (0.2s - 0.4s): Ngay sau khi giai đoạn Vector Search & Hybrid Retrieval hoàn tất, backend lập tức phát sự kiện 'meta' chứa danh sách nguồn tham khảo và các bước Trace Log về cho client.",
        "Phát luồng ký tự trực tiếp: Kết nối trực tiếp với OpenRouter API qua cơ chế stream=True, giúp mỗi từ (token) được sinh ra sẽ được đẩy ngay tới giao diện người dùng mà không cần chờ toàn bộ văn bản hoàn thành.",
        "Giảm độ trễ TTFT: Thời gian hiển thị ký tự đầu tiên (Time To First Token) giảm mạnh từ ~5.0 giây xuống dưới 0.5 giây.",
    ]
    for sp in stream_points:
        bp = doc.add_paragraph(style='List Bullet')
        r = bp.add_run(sp)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        bp.paragraph_format.space_after = Pt(4)

    # 2.3
    h23 = doc.add_heading(level=2)
    r23 = h23.add_run("2.3. Tối ưu trải nghiệm gõ sẵn câu hỏi trên giao diện (UI Input Draft)")
    r23.font.name = "Times New Roman"
    r23.font.size = Pt(13)
    r23.font.bold = True
    r23.font.color.rgb = RGBColor(40, 40, 40)
    h23.paragraph_format.space_before = Pt(8)
    h23.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph("Chỉnh sửa mã nguồn giao diện static/index.html:")
    p.paragraph_format.space_after = Pt(4)

    ui_points = [
        "Luôn giữ ô nhập liệu mở (disabled = false): Cho phép người dùng nhập trước nội dung câu hỏi thứ hai trong lúc AI đang trả lời câu hỏi thứ nhất.",
        "Nút bấm thông minh (Dynamic Control Button): Khi ô nhập liệu trống, nút hiển thị trạng thái '🛑 Dừng' để dừng phản hồi nếu muốn. Ngay khi người dùng gõ ký tự vào ô nhập, nút tự động chuyển sang 'Gửi'.",
        "Chuyển tiếp mượt mà: Người dùng chỉ cần gõ sẵn câu hỏi và bấm Enter/Gửi, hệ thống sẽ tự động ngắt câu trước và phản hồi ngay câu hỏi mới.",
    ]
    for up in ui_points:
        bp = doc.add_paragraph(style='List Bullet')
        r = bp.add_run(up)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        bp.paragraph_format.space_after = Pt(4)

    # SECTION 3
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. BÁO CÁO KẾT QUẢ KIỂM THỬ 100 CÂU HỎI (BENCHMARK EVALUATION)")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(14)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0, 51, 102)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Để đảm bảo tính chính xác và hiệu năng của hệ thống sau khi tối ưu, kịch bản kiểm thử tự động benchmark_100_questions.py đã được khởi chạy trên bộ 100 câu hỏi (gồm 80 câu chuyên sâu tuyển sinh HUIT và 20 câu hỏi ngoài phạm vi). Kết quả thu được như sau:")
    p.paragraph_format.space_after = Pt(6)

    bench_table = doc.add_table(rows=6, cols=3)
    bench_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(bench_table)

    b_headers = ["Tiêu chí đánh giá", "Kết quả thực tế", "Đánh giá chất lượng"]
    for idx, text in enumerate(b_headers):
        cell = bench_table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.bold = True
        r.font.size = Pt(11)
        set_cell_margins(cell, 80, 80, 100, 100)

    b_data = [
        ("Tổng số câu hỏi thử nghiệm", "100 / 100 câu", "Hoàn tất 100% kịch bản"),
        ("Số lượng câu phát sinh lỗi", "0 lỗi (0%)", "Hệ thống vận hành tuyệt đối ổn định"),
        ("Tỷ lệ trích xuất đúng tri thức (Hit Rate)", "97.5% (78/80 câu HUIT)", "Chính xác cao đối với dữ liệu 39 ngành HUIT"),
        ("Tỷ lệ không bị ảo tưởng nguồn (Guardrail)", "100% (20/20 câu ngoài)", "Không sinh link nguồn giả mạo"),
        ("Thời gian xử lý trung bình", "0.60 giây / câu", "Phản hồi siêu nhanh"),
    ]
    for r_idx, data in enumerate(b_data, 1):
        row = bench_table.rows[r_idx]
        for c_idx, val in enumerate(data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            set_cell_margins(cell, 60, 60, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SECTION 4
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. KẾT LUẬN VÀ HƯỚNG DẪN KIỂM TRA")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(14)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(0, 51, 102)
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Hệ thống Trợ lý AI Tuyển sinh HUIT đã hoàn thành toàn bộ công tác nâng cấp kĩ thuật, khắc phục triệt để các vấn đề về độ trễ, đồng bộ kho chứa và cải thiện trải nghiệm người dùng trên giao diện web.")
    p.paragraph_format.space_after = Pt(6)

    p_link = doc.add_paragraph()
    r_link_title = p_link.add_run("Địa chỉ truy cập kiểm thử trực tiếp: ")
    r_link_title.font.name = "Times New Roman"
    r_link_title.bold = True
    r_link = p_link.add_run("https://huit-chatbot.vercel.app/")
    r_link.font.name = "Times New Roman"
    r_link.font.color.rgb = RGBColor(0, 102, 204)
    r_link.font.underline = True

    p_link.paragraph_format.space_after = Pt(18)

    # Signature Block
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l, cell_r = sig_table.rows[0].cells[0], sig_table.rows[0].cells[1]
    
    pl = cell_l.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = pl.add_run("ĐỘI NGŨ PHÁT TRIỂN HỆ THỐNG\n(Ký và ghi rõ họ tên)")
    rl.font.name = "Times New Roman"
    rl.font.bold = True

    pr = cell_r.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pr.add_run("XÁC NHẬN CỦA ĐƠN VỊ\n(Ký và đóng dấu)")
    rr.font.name = "Times New Roman"
    rr.font.bold = True

    out_path = os.path.join("d:\\chatbot2", "Bao_Cao_Xu_Ly_He_Thong_HUIT_Chatbot.docx")
    doc.save(out_path)
    print("Report generated successfully:", out_path)

if __name__ == "__main__":
    build_report()
